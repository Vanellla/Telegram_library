from docx import Document
from docx.shared import Mm
import os
from qr import make_qr_and_return_path, qrcodes
from pathlib import Path
import os

from settings import doc_dir
from datetime import datetime

def make_word_table_with_qr(sp: list, file_name: str) -> str:  # ['Полина', 'Варя', 'Паша', 'Сергей']
    doc = Document()  # создаем пустой док
    cols = 5  # число столбцов в таблице
    # вычисление количества строк таблицы
    row = len(sp) // cols
    if len(sp) % cols != 0:
        row += 1

    table = doc.add_table(rows=row, cols=cols)  # Добавление таблицы в документ
    table.style = 'Table Grid'  # задаем рамку
    for i, code in enumerate(sp):
        # i - порядковый номер кода из списка
        # code - текст, который будет помещен в qr код
        row_index = i // 5  # вычисляем строку
        cell_index = i % 5  # вычисляем столбец
        # получаем доступ к содержимому ячейки
        paragr = table.rows[row_index].cells[cell_index].paragraphs[0]
        run = paragr.add_run()
        # создаем qr-код и получаем путь до него на диске
        qr_img = make_qr_and_return_path(code)
        # вставляем код в ячейку
        run.add_picture(qr_img, width=Mm(25))
        os.remove(qr_img)
    if not (os.path.isdir(doc_dir)):
        os.mkdir(doc_dir)
    date = datetime.now().strftime("%H_%M_%S__%d_%m")
    doc_path = os.path.join(doc_dir, f'{file_name}_{date}.docx')
    doc.save(doc_path)  # сохраняем документ
    return doc_path
